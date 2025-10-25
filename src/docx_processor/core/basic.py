"""
Basic DOCX processor using pure python-docx.
This is the foundation that works without any external dependencies.
"""

import logging
import json
from collections import defaultdict
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from io import BytesIO

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from PIL import Image
import re
import html

from ..models import (
    ProcessingResult, ProcessingConfig, ProcessingMode,
    SectionInfo, ImageInfo, TableInfo, TOCEntry
)
from ..exceptions import ProcessingError
from ..utils.text_utils import clean_string, remove_section_headers


class BasicProcessor:
    """
    Pure python-docx processor for basic DOCX document processing.
    
    This processor handles:
    - Document structure extraction
    - Text content extraction with hierarchy
    - Table extraction and HTML conversion
    - Basic image extraction
    - Table of contents generation
    """
    
    def __init__(self, config: ProcessingConfig, logger: logging.Logger):
        """
        Initialize the basic processor.
        
        Args:
            config: Processing configuration
            logger: Logger instance
        """
        self.config = config
        self.logger = logger
    
    def process(self, file_path: Path, config: ProcessingConfig) -> ProcessingResult:
        """
        Process a DOCX document using basic python-docx functionality.
        
        Args:
            file_path: Path to the DOCX file
            config: Processing configuration
            
        Returns:
            ProcessingResult with extracted content
        """
        self.logger.info(f"Starting basic processing of {file_path}")
        
        try:
            # Load the document
            doc = Document(file_path)
            
            # Extract structured content
            content_dict, content_without_children, image_sections, table_mapping, table_dict, toc = self._process_document_structure(doc)
            
            # Extract images
            images_dict = {}
            if config.save_images:
                images_dict = self._extract_images(doc, image_sections, config.output_dir)
            
            # Process tables
            tables_dict = {}
            if config.save_tables:
                tables_dict = self._process_tables(table_dict, config.output_dir)
            
            # Extract headers and footers
            headers_footers = {}
            if config.include_headers_footers:
                headers_footers = self._extract_headers_footers(doc)
            
            # Extract endnotes
            endnotes = self._extract_endnotes(doc)
            
            # Build TOC if requested
            toc_entries = None
            if config.extract_toc and toc:
                toc_entries = self._build_toc_entries(toc)
            
            # Create section info hierarchy
            content_hierarchy = self._build_content_hierarchy(content_dict)
            
            # Create result
            result = ProcessingResult(
                content=content_dict,
                content_without_children=content_without_children,
                content_hierarchy=content_hierarchy,
                images=images_dict,
                tables=tables_dict,
                toc=toc_entries,
                headers_footers=headers_footers,
                endnotes=endnotes,
                processing_mode=ProcessingMode.BASIC
            )
            
            # Save outputs if output directory is specified
            if config.output_dir:
                self._save_outputs(result, config.output_dir, config)
            
            self.logger.info(f"Basic processing completed. Extracted {len(content_dict)} sections, {len(images_dict)} images, {len(tables_dict)} tables")
            return result
            
        except Exception as e:
            self.logger.error(f"Basic processing failed: {e}")
            raise ProcessingError(f"Failed to process document: {e}") from e
    
    def _process_document_structure(self, doc: Document) -> Tuple[Dict, Dict, List, Dict, Dict, Any]:
        """
        Extract the document structure with proper hierarchy.
        
        This is adapted from the original DOCXProcessor.process_docx method.
        """
        content_dict = {}
        counters = {}
        table_mapping = {}
        table_dict = {}
        
        stack = []
        img_title = []
        title = None
        current_content = None
        table_count = 1
        
        self.logger.debug("Processing document structure with hierarchy")
        
        # Process paragraphs and tables together to preserve order
        # We'll iterate through the document elements in order
        for paragraph in doc.paragraphs:
            content = paragraph.text
            
            # Skip empty paragraphs without images
            image_is_present = self._has_image_in_paragraph(paragraph)
            if content.strip() == '' and not image_is_present:
                continue
            
            heading_level = self._get_heading_level(paragraph)
            
            if heading_level:
                # Initialize or reset counters for deeper levels
                if heading_level not in counters:
                    counters[heading_level] = 0
                counters[heading_level] += 1
                
                # Reset counters for sublevels
                for level in list(counters.keys()):
                    if level > heading_level:
                        counters[level] = 0
                
                # Generate the numbering for the current heading
                numbering = '.'.join(str(counters[level]) for level in sorted(counters) if counters[level] > 0)
                
                # If we are encountering a new heading level, save the previous content
                if current_content is not None:
                    content_dict[current_content["title"]] = current_content
                
                # Determine the parent
                parent = None
                if stack:
                    if stack[-1]["level"] < heading_level:
                        parent = stack[-1]["title"]
                    elif stack[-1]["level"] == heading_level:
                        parent = stack[-1]['parent']
                    else:
                        while stack and stack[-1]["level"] >= heading_level:
                            stack.pop()
                        parent = stack[-1]["title"] if stack else None
                
                # Create title with numbering if needed
                pattern = r'^\s*\d+(\.\d+)*\.?\s*'
                if re.match(pattern, paragraph.text):
                    title = f"{paragraph.text}"
                else:
                    title = f"{numbering} {paragraph.text}"
                
                if title.strip() == '':
                    title = 'Orphaned Section'
                
                current_content = {
                    "title": title,
                    "level": heading_level,
                    "content": "",
                    "parent": parent,
                    "children": []
                }
                
                # Manage the stack to handle hierarchy
                while stack and stack[-1]["level"] >= heading_level:
                    stack.pop()
                
                if parent and parent in content_dict:
                    content_dict[parent]["children"].append(current_content["title"])
                
                stack.append({"title": current_content["title"], "level": heading_level, "parent": parent})
                
            else:
                # Append content to the current heading
                if current_content is not None:
                    current_content["content"] += content + "\n"
                else:
                    # If we haven't encountered any heading yet, create a title page
                    current_content = {
                        "title": 'TITLE PAGE',
                        "level": 1,
                        "content": content + "\n",
                        "parent": None,
                        "children": []
                    }
            
            # Create a dict for the sections that contain images
            if image_is_present:
                current_title = current_content["title"] if current_content else 'TITLE PAGE'
                if current_title not in img_title:
                    img_title.append(current_title)
        
        # Process tables (simplified for now - we'll enhance this later)
        for table in doc.tables:
            table_html = self._table_to_html(table)
            table_file_name = f"table_{table_count}.html"
            table_dict[table_file_name] = table_html
            
            # Add table to current section
            if current_content is not None:
                current_content["content"] += f"Table\n{table_html}\n"
                
                # Track table mapping
                section_title = current_content["title"]
                if section_title in table_mapping:
                    table_mapping[section_title].append(table_file_name)
                else:
                    table_mapping[section_title] = [table_file_name]
            
            table_count += 1
        
        # Append the last content item
        if current_content is not None:
            content_dict[current_content["title"]] = current_content
        
        # Create content without children (individual sections only)
        final_dict_without_child = {item["title"]: item["content"] for item in content_dict.values()}
        
        # Build table of contents (simplified for now)
        toc = self._build_toc(list(content_dict.values()))
        
        # Recursive function to aggregate content from children to parent
        def aggregate_content(title: str) -> None:
            item = content_dict[title]
            for child_title in item["children"]:
                if child_title in content_dict:  # Safety check
                    aggregate_content(child_title)
                    item["content"] += f"\n\n{content_dict[child_title]['title']}\n\n{content_dict[child_title]['content']}"
        
        # Start aggregation from the top-level items
        top_level_titles = [title for title, item in content_dict.items() if item["parent"] is None]
        for title in top_level_titles:
            aggregate_content(title)
        
        # Prepare the final dictionary with title as key and combined content as value
        final_dict = {item["title"]: item["content"] for item in content_dict.values()}
        
        self.logger.debug(f"Extracted {len(final_dict)} sections with hierarchy")
        
        return final_dict, final_dict_without_child, img_title, table_mapping, table_dict, toc
    
    def _has_image_in_paragraph(self, paragraph: Paragraph) -> bool:
        """Check if a paragraph contains an image."""
        try:
            return (
                'w:drawing' in paragraph._p.xml and (
                    ("graphicData" in paragraph._p.xml) or
                    ("wp:anchor" in paragraph._p.xml)
                )
                or ("imagedata" in paragraph._p.xml)
            )
        except Exception:
            return False
    
    def _get_heading_level(self, paragraph: Paragraph) -> Optional[int]:
        """Get the heading level of a paragraph."""
        try:
            if paragraph.style.name.startswith('Heading'):
                try:
                    return int(paragraph.style.name.split()[-1])
                except:
                    # Handle ill-formatted heading styles
                    return int(list(set(re.findall(r'\d+', paragraph.style.name)))[0])
            return None
        except Exception:
            return None
    
    def _table_to_html(self, table: Table) -> str:
        """
        Convert a DOCX table to HTML representation with advanced styling.
        Based on the original DOCXProcessor.table_to_html method.
        """
        try:
            def rgb_to_hex(rgb: Tuple[int, int, int]) -> str:
                if rgb is None:
                    return None
                return '#{:02x}{:02x}{:02x}'.format(rgb[0], rgb[1], rgb[2])

            def get_cell_style(cell: _Cell) -> str:
                style = "border: 1px solid black; padding: 5px;"
                if cell.vertical_alignment:
                    style += f"vertical-align:{cell.vertical_alignment};"
                if cell._tc.tcPr.tcW is not None:
                    width = cell._tc.tcPr.tcW.w
                    if width is not None:
                        style += f"width:{width / 15}pt;"
                if cell.paragraphs and cell.paragraphs[0].runs:
                    run = cell.paragraphs[0].runs[0]
                    if run.font.color.rgb is not None:
                        color = rgb_to_hex(run.font.color.rgb)
                        if color:
                            style += f"color:{color};"
                    if run.font.size is not None:
                        style += f"font-size:{run.font.size.pt}pt;"
                    if run.font.bold:
                        style += "font-weight:bold;"
                    if run.font.italic:
                        style += "font-style:italic;"
                return style

            def get_paragraph_style(paragraph: Paragraph) -> str:
                style = ""
                if paragraph.alignment:
                    style += f"text-align:{paragraph.alignment};"
                if paragraph.style:
                    if paragraph.style.font.color.rgb is not None:
                        color = rgb_to_hex(paragraph.style.font.color.rgb)
                        if color:
                            style += f"color:{color};"
                    if paragraph.style.font.size is not None:
                        style += f"font-size:{paragraph.style.font.size.pt}pt;"
                return style

            html_output = "<table style='border-collapse: collapse; border: 1px solid black;'>"

            # Table header
            if table.rows:
                html_output += "<thead><tr>"
                for cell in table.rows[0].cells:
                    cell_style = get_cell_style(cell)
                    html_output += f"<th style='{cell_style}'>"
                    for para in cell.paragraphs:
                        para_style = get_paragraph_style(para)
                        html_output += f"<p style='{para_style}'>{html.escape(para.text)}</p>"
                    html_output += "</th>"
                html_output += "</tr></thead>"

            # Table body
            html_output += "<tbody>"
            for row in table.rows[1:]:  # Start from the second row
                html_output += "<tr>"
                for cell in row.cells:
                    cell_style = get_cell_style(cell)
                    html_output += f"<td style='{cell_style}'>"
                    for para in cell.paragraphs:
                        para_style = get_paragraph_style(para)
                        if para.style and 'Bullets' in para.style.name:
                            html_output += f"<ul style='{para_style}'><li>{html.escape(para.text)}</li></ul>"
                        else:
                            html_output += f"<p style='{para_style}'>{html.escape(para.text)}</p>"
                    html_output += "</td>"
                html_output += "</tr>"
            html_output += "</tbody>"
            html_output += "</table>"
            
            return html_output   
        except Exception as e:
            self.logger.error(f"Error converting table to HTML: {e}")
            return f"<p>Error processing table: {e}</p>"
    
    def _build_toc(self, sections: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Build a table of contents from section hierarchy."""
        toc = {}
        
        # Sort sections by level to build hierarchy
        sorted_sections = sorted(sections, key=lambda x: x.get('level', 1))
        
        for section in sorted_sections:
            title = section['title']
            level = section.get('level', 1)
            parent = section.get('parent')
            
            toc[title] = {
                'level': level,
                'parent': parent,
                'children': section.get('children', [])
            }
        
        return toc
    
    def _extract_images(self, doc: Document, image_sections: List[str], output_dir: Optional[Path]) -> Dict[str, ImageInfo]:
        """
        Extract images from the document with advanced processing.
        Based on the original DOCXProcessor.extract_images method.
        """
        images_dict = {}
        section_image_mapping = {}
        image_count = 1
        
        try:
            # Extract images from different formats: blips, VML, flowcharts
            for element in doc.element.body.iter():
                if element.tag == qn('w:drawing'):
                    for child in element.iter():
                        if child.tag == qn('a:blip'):
                            image_info = self._extract_blip_image(child, doc, output_dir, image_count)
                            if image_info:
                                images_dict[image_info.filename] = image_info
                                image_count += 1
                        
                        elif child.tag == qn("wp:anchor"):
                            # Handle flowchart images (simplified - would need external processor)
                            self.logger.debug("Flowchart image detected but skipped in basic mode")
                
                # Handle VML images
                if element.tag == "{urn:schemas-microsoft-com:vml}imagedata":
                    image_info = self._extract_vml_image(element, doc, output_dir, image_count)
                    if image_info:
                        images_dict[image_info.filename] = image_info
                        image_count += 1
                        
            # Map images to their sections based on order
            section_index = 0
            for filename, image_info in images_dict.items():
                if section_index < len(image_sections):
                    section_title = image_sections[section_index]
                    if section_title not in section_image_mapping:
                        section_image_mapping[section_title] = []
                    section_image_mapping[section_title].append(filename)
                    # Update image_info with section mapping
                    image_info.section = section_title
                    section_index += 1
                    
            self.logger.info(f"Successfully extracted {len(images_dict)} images from document")
            
        except Exception as e:
            self.logger.error(f"Error extracting images: {e}")
        
        return images_dict
    
    def _extract_blip_image(self, blip_child, doc: Document, output_dir: Optional[Path], image_count: int) -> Optional[ImageInfo]:
        """
        Extract a single blip image with advanced processing including cropping.
        Based on the original DOCXProcessor._extract_blips method.
        """
        try:
            rId = blip_child.attrib[qn('r:embed')]
            image_part = doc.part.related_parts[rId]
            
            # Get image format
            content_type = image_part.content_type
            file_extension = content_type.split('/')[-1] if '/' in content_type else 'png'
            
            # Generate filename
            image_filename = f"image_{image_count}.{file_extension}"
            
            # Get original image bytes
            image_bytes = image_part.blob
            
            # Apply auto-cropping if srcRect is present
            processed_image_bytes = self._apply_auto_cropping(blip_child, image_bytes)
            
            # Save image if output directory is provided
            if output_dir:
                images_dir = output_dir / "images"
                images_dir.mkdir(parents=True, exist_ok=True)
                image_path = images_dir / image_filename
                with open(image_path, "wb") as f:
                    f.write(processed_image_bytes)
            
            # Get image dimensions
            try:
                img = Image.open(BytesIO(processed_image_bytes))
                width, height = img.size
            except:
                width, height = None, None
            
            return ImageInfo(
                filename=image_filename,
                size_bytes=len(processed_image_bytes),
                width=width,
                height=height,
                format=file_extension.upper()
            )
            
        except KeyError:
            self.logger.error(f"Image with rId {rId} not found.")
            return None
        except Exception as e:
            self.logger.error(f"Error extracting blip image: {e}")
            return None
    
    def _apply_auto_cropping(self, blip_element, image_bytes: bytes) -> bytes:
        """
        Apply cropping to an image based on the srcRect information in the DOCX element.
        
        The srcRect element defines cropping as percentages in thousandths of a percent:
        - l: left crop percentage (distance from left edge)
        - t: top crop percentage (distance from top edge)  
        - r: right crop percentage (distance from right edge)
        - b: bottom crop percentage (distance from bottom edge)
        """
        try:
            # Look for srcRect element in the parent blipFill element
            blip_fill = blip_element.getparent()
            if blip_fill is None:
                return image_bytes
                
            # Find srcRect element within the blipFill
            src_rect = None
            for child in blip_fill:
                if child.tag == qn('a:srcRect'):
                    src_rect = child
                    break
            
            if src_rect is None:
                return image_bytes
                
            # Extract crop percentages from srcRect attributes
            # Values are in thousandths of a percent (divide by 1000 to get percentage)
            left_crop = float(src_rect.get('l', '0')) / 1000.0
            top_crop = float(src_rect.get('t', '0')) / 1000.0
            right_crop = float(src_rect.get('r', '0')) / 1000.0
            bottom_crop = float(src_rect.get('b', '0')) / 1000.0
            
            # If no cropping is specified, return original image
            if left_crop == 0 and top_crop == 0 and right_crop == 0 and bottom_crop == 0:
                return image_bytes
                
            self.logger.info(f"Applying srcRect cropping: left={left_crop}%, top={top_crop}%, right={right_crop}%, bottom={bottom_crop}%")
            
            # Load image using PIL
            image = Image.open(BytesIO(image_bytes))
            original_width, original_height = image.size
            
            # Calculate crop coordinates in pixels
            left_pixels = int((left_crop / 100.0) * original_width)
            top_pixels = int((top_crop / 100.0) * original_height)
            right_pixels = original_width - int((right_crop / 100.0) * original_width)
            bottom_pixels = original_height - int((bottom_crop / 100.0) * original_height)
            
            # Ensure coordinates are within image bounds
            left_pixels = max(0, min(left_pixels, original_width))
            top_pixels = max(0, min(top_pixels, original_height))
            right_pixels = max(left_pixels, min(right_pixels, original_width))
            bottom_pixels = max(top_pixels, min(bottom_pixels, original_height))
            
            self.logger.debug(f"Cropping image from ({left_pixels}, {top_pixels}) to ({right_pixels}, {bottom_pixels})")
            
            # Crop the image
            cropped_image = image.crop((left_pixels, top_pixels, right_pixels, bottom_pixels))
            
            # Convert back to bytes
            output_buffer = BytesIO()
            format = image.format if image.format else 'PNG'
            cropped_image.save(output_buffer, format=format)
            cropped_bytes = output_buffer.getvalue()
            
            self.logger.info(f"Successfully cropped image from {original_width}x{original_height} to {cropped_image.size[0]}x{cropped_image.size[1]}")
            
            return cropped_bytes
            
        except Exception as e:
            self.logger.error(f"Error applying srcRect cropping: {e}")
            return image_bytes
    
    def _extract_vml_image(self, vml_element, doc: Document, output_dir: Optional[Path], image_count: int) -> Optional[ImageInfo]:
        """
        Extract VML image data from a DOCX element.
        Simplified version - full VML processing would require LibreOffice conversion.
        """
        try:
            rId = vml_element.attrib[qn("r:id")]
            image_part = doc.part.related_parts[rId]
            
            if image_part.content_type == "image/x-emf":
                # EMF images would need conversion (LibreOffice or similar)
                # For now, we'll skip EMF images in basic mode
                self.logger.debug("EMF image detected but skipped in basic mode (requires LibreOffice conversion)")
                return None
            else:
                # Handle regular image formats
                file_extension = image_part.content_type.split('/')[-1] if '/' in image_part.content_type else 'png'
                image_filename = f"image_{image_count}.{file_extension}"
                image_bytes = image_part.blob
                
                # Save image if output directory is provided
                if output_dir:
                    images_dir = output_dir / "images"
                    images_dir.mkdir(parents=True, exist_ok=True)
                    image_path = images_dir / image_filename
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)
                
                # Get image dimensions
                try:
                    img = Image.open(BytesIO(image_bytes))
                    width, height = img.size
                except:
                    width, height = None, None
                
                return ImageInfo(
                    filename=image_filename,
                    size_bytes=len(image_bytes),
                    width=width,
                    height=height,
                    format=file_extension.upper()
                )
                
        except Exception as e:
            self.logger.error(f"Error extracting VML image: {e}")
            return None
    
    def _extract_headers_footers(self, doc: Document) -> Dict[str, Any]:
        """
        Extract headers and footers from all sections of the document.
        Based on the original DOCXProcessor.extract_header_footer_content method.
        """
        headers_footers = {
            "headers": {},
            "footers": {},
            "unique_headers": set(),
            "unique_footers": set()
        }
        
        try:
            for section_idx, section in enumerate(doc.sections):
                section_name = f"section_{section_idx + 1}"
                
                # Extract header content
                header_content = ""
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            header_content += paragraph.text + "\n"
                
                # Extract footer content  
                footer_content = ""
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            footer_content += paragraph.text + "\n"
                
                # Store section-specific content
                headers_footers["headers"][section_name] = header_content.strip()
                headers_footers["footers"][section_name] = footer_content.strip()
                
                # Track unique content to avoid repetition
                if header_content.strip():
                    headers_footers["unique_headers"].add(header_content.strip())
                if footer_content.strip():
                    headers_footers["unique_footers"].add(footer_content.strip())
            
            # Convert sets to lists for JSON serialization
            headers_footers["unique_headers"] = list(headers_footers["unique_headers"])
            headers_footers["unique_footers"] = list(headers_footers["unique_footers"])
            
            self.logger.debug(f"Extracted headers from {len(headers_footers['headers'])} sections")
            self.logger.debug(f"Found {len(headers_footers['unique_headers'])} unique headers")
            self.logger.debug(f"Found {len(headers_footers['unique_footers'])} unique footers")
            
        except Exception as e:
            self.logger.error(f"Error extracting headers/footers: {e}")
        
        return headers_footers
    
    def _extract_endnotes(self, doc: Document) -> Dict[str, str]:
        """
        Extract endnotes from the document.
        Based on the original DOCXProcessor.get_all_endnotes method.
        """
        endnotes = {}
        endnotes_text = ""
        
        try:
            # Access the document's XML to find endnotes
            from docx.oxml.ns import qn
            
            # Look for endnotes in the document part
            if hasattr(doc.part, 'package') and doc.part.package:
                for part in doc.part.package.parts:
                    if 'endnotes' in part.partname:
                        # Parse endnotes XML
                        endnotes_xml = part.blob
                        if endnotes_xml:
                            endnotes_text = self._parse_endnotes_xml(endnotes_xml)
                            break
            
            # If we found endnotes, organize them
            if endnotes_text:
                endnotes["raw_endnotes"] = endnotes_text
                endnotes["formatted_endnotes"] = self._format_endnotes(endnotes_text)
                self.logger.debug(f"Extracted endnotes content: {len(endnotes_text)} characters")
            else:
                self.logger.debug("No endnotes found in document")
            
        except Exception as e:
            self.logger.error(f"Error extracting endnotes: {e}")
        
        return endnotes
    
    def _parse_endnotes_xml(self, endnotes_xml: bytes) -> str:
        """Parse endnotes XML content to extract text."""
        try:
            import xml.etree.ElementTree as ET
            root = ET.fromstring(endnotes_xml)
            
            # Extract text from all text elements
            text_content = []
            for text_elem in root.iter():
                if text_elem.tag.endswith('}t') and text_elem.text:  # w:t elements
                    text_content.append(text_elem.text)
            
            return '\n'.join(text_content)
        except Exception as e:
            self.logger.error(f"Error parsing endnotes XML: {e}")
            return ""
    
    def _format_endnotes(self, endnotes_text: str) -> str:
        """Format endnotes text for better readability."""
        if not endnotes_text:
            return ""
        
        # Basic formatting - could be enhanced based on needs
        lines = endnotes_text.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                # Add some basic formatting for numbered references
                if line and line[0].isdigit():
                    formatted_lines.append(f"\n{line}")
                else:
                    formatted_lines.append(line)
        
        return '\n'.join(formatted_lines).strip()
    
    def _process_tables(self, table_dict: Dict[str, str], output_dir: Optional[Path]) -> Dict[str, TableInfo]:
        """Process and save tables."""
        tables_info = {}
        
        if output_dir:
            tables_dir = output_dir / "tables"
            tables_dir.mkdir(parents=True, exist_ok=True)
            
            for filename, html_content in table_dict.items():
                table_path = tables_dir / filename
                with open(table_path, "w", encoding="utf-8") as f:
                    f.write(html_content)
                
                # Create table info
                tables_info[filename] = TableInfo(
                    filename=filename,
                    rows=html_content.count("<tr>"),
                    columns=html_content.count("<th>") or html_content.count("<td>")
                )
        
        return tables_info
    
    def _build_toc_entries(self, toc: Any) -> List[TOCEntry]:
        """Convert TOC structure to TOCEntry objects."""
        # This would need to be implemented based on your TOC structure
        return []
    
    def _build_content_hierarchy(self, content_dict: Dict) -> Dict[str, SectionInfo]:
        """Build content hierarchy from the content dictionary."""
        hierarchy = {}
        
        for title, content in content_dict.items():
            # For now, create simple hierarchy since our content_dict is just title->content
            hierarchy[title] = SectionInfo(
                title=title,
                content=content,
                level=1,  # Default level
                parent=None,
                children=[]
            )
        
        return hierarchy
    
    def _build_toc(self, elements: List[Dict]) -> Any:
        """Build table of contents from document elements."""
        # Simplified version - you can adapt your original build_toc method here
        return None
    
    def _save_outputs(self, result: ProcessingResult, output_dir: Path, config: ProcessingConfig) -> None:
        """Save processing outputs to files."""
        if config.save_content:
            # Save content as JSON
            content_path = output_dir / "content.json"
            with open(content_path, "w", encoding="utf-8") as f:
                json.dump(result.content, f, indent=2, ensure_ascii=False)
            
            if result.content_without_children:
                content_without_children_path = output_dir / "content_without_children.json"
                with open(content_without_children_path, "w", encoding="utf-8") as f:
                    json.dump(result.content_without_children, f, indent=2, ensure_ascii=False)
            
            # Save headers and footers
            if result.headers_footers:
                headers_footers_path = output_dir / "headers_footers.json"
                with open(headers_footers_path, "w", encoding="utf-8") as f:
                    json.dump(result.headers_footers, f, indent=2, ensure_ascii=False)
            
            # Save endnotes
            if result.endnotes:
                endnotes_path = output_dir / "endnotes.json"
                with open(endnotes_path, "w", encoding="utf-8") as f:
                    json.dump(result.endnotes, f, indent=2, ensure_ascii=False)
            
            # Save HTML content (enhanced mode)
            if result.html_content:
                html_path = output_dir / "content.html"
                with open(html_path, "w", encoding="utf-8") as f:
                    f.write(result.html_content)
        
        # Update result with output paths
        result.output_paths = {
            "content": output_dir / "content.json",
            "images": output_dir / "images",
            "tables": output_dir / "tables",
            "headers_footers": output_dir / "headers_footers.json",
            "endnotes": output_dir / "endnotes.json"
        }
        
        # Add enhanced mode output paths if available
        if result.html_content:
            result.output_paths["html"] = output_dir / "content.html"
        if result.page_screenshots:
            result.output_paths["pages"] = output_dir / "pages"